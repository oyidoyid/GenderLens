import os
import re
import io
import requests
import time

from flask import (
    Flask, render_template, request, redirect,
    url_for, session, send_file, flash
)
from flask_session import Session
from datetime import timedelta
import docx  # pip install python-docx
from authlib.integrations.flask_client import OAuth

from biased_words import highlight_text, calculate_bias_percentage
from certificate import generate_pdf_bytes
from werkzeug.exceptions import RequestEntityTooLarge


app = Flask(__name__)
app.secret_key = "CHANGE_THIS_SECRET_KEY"   
app.permanent_session_lifetime = timedelta(days=7)

# --- START OF NEW SESSION CODE ---
# Create folder for session files
SESSION_DIR = os.path.join(os.getcwd(), 'flask_session')
if not os.path.exists(SESSION_DIR):
    os.makedirs(SESSION_DIR)

# Configure Server-Side Session
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = SESSION_DIR
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True

# Initialize the extension
Session(app)
# --- END OF NEW SESSION CODE ---


app.config["MAX_CONTENT_LENGTH"] = 50 *1024 * 1024  


app.config["GOOGLE_CLIENT_ID"] = os.environ.get("GOOGLE_CLIENT_ID")
app.config["GOOGLE_CLIENT_SECRET"] = os.environ.get("GOOGLE_CLIENT_SECRET")
app.config["GOOGLE_DISCOVERY_URL"] = (
    "https://accounts.google.com/.well-known/openid-configuration"
)

if not app.config["GOOGLE_CLIENT_ID"] or not app.config["GOOGLE_CLIENT_SECRET"]:
    raise RuntimeError(
        "GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET must be set as environment variables."
    )

oauth = OAuth(app)
google = oauth.register(
    name="google",
    client_id=app.config["GOOGLE_CLIENT_ID"],
    client_secret=app.config["GOOGLE_CLIENT_SECRET"],
    server_metadata_url=app.config["GOOGLE_DISCOVERY_URL"],
    client_kwargs={"scope": "openid email profile"}
)


USAGE_LIMITS = {
    "guest": 3,         # not logged in
    "user": 30,         # logged in with Google
    "subscriber": 1000  # effectively unlimited
}


def get_role():
    """Return current user's role (guest / user / subscriber)."""
    return session.get("user_role", "guest")


def is_logged_in():
    """Return True if user_email is in session."""
    return session.get("user_email") is not None


def get_usage_count():
    """How many analyses the current session has performed."""
    return session.get("usage_count", 0)


def increment_usage():
    """Increase the analysis counter for this session."""
    session["usage_count"] = get_usage_count() + 1


def get_max_uses_for_current_user():
    """Maximum allowed analyses for the current role."""
    role = get_role()
    return USAGE_LIMITS.get(role, USAGE_LIMITS["guest"])


# -------------------------------------------------
# Helper: apply suggestions to text
# -------------------------------------------------
def apply_suggestions_to_text(original_text, suggestions_list):
    """
    Replace biased terms in original_text with their first suggested replacement.
    suggestions_list: list of dicts like {"word": "chairman", "replacements": ["chair", "chairperson"]}
    """
    cleaned_text = original_text

    # Replace longer phrases first
    sorted_suggestions = sorted(
        suggestions_list,
        key=lambda x: len(x["word"]),
        reverse=True
    )

    for item in sorted_suggestions:
        biased_term = item["word"]
        replacements = item["replacements"]
        if not replacements:
            continue
        replacement = replacements[0]

        pattern = re.compile(
            r"(?<!\w)" + re.escape(biased_term) + r"(?!\w)",
            flags=re.IGNORECASE
        )
        cleaned_text = pattern.sub(replacement, cleaned_text)

    return cleaned_text



@app.route("/", methods=["GET", "POST"])
@app.route("/dashboard", methods=["GET", "POST"])
def dashboard():
    role = get_role()
    max_uses = get_max_uses_for_current_user()
    current_usage = get_usage_count()

   
    if current_usage >= max_uses and role != "subscriber":
        flash(f"Usage limit reached for your account type ({role}). Please log in for more analyses.")
        remaining_uses = 0
        usage_count = current_usage
        last_response_time_ms = session.get("last_response_time_ms")

        return render_template(
            "dashboard.html",
            text="",
            result=None,
            suggestions=[],
            logged_in=is_logged_in(),
            user_email=session.get("user_email"),
            remaining_uses=remaining_uses,
            role=role,
            usage_count=usage_count,
            max_guest_uses=max_uses,
            last_response_time_ms=last_response_time_ms,
        )


    text = ""
    result = None
    suggestions = []
    from_file = False

    if request.method == "POST":
       
        text = request.form.get("text_input", "").strip()

        # 2) Optional: file upload (.docx)
        uploaded_file = request.files.get("file_input")
        if uploaded_file and uploaded_file.filename:
            # NEW: defensive size check for 100 MB (in addition to MAX_CONTENT_LENGTH)
            uploaded_file.seek(0, os.SEEK_END)
            file_size = uploaded_file.tell()
            uploaded_file.seek(0)

            max_size_bytes = 100 * 1024 * 1024
            if file_size > max_size_bytes:
                flash("Uploaded file is too large. Maximum allowed size is 100 MB.")
                # Just show form again; do not process this file
                remaining_uses = None
                if role != "subscriber":
                    remaining_uses = max(0, max_uses - get_usage_count())
                usage_count = get_usage_count()
                max_uses_local = get_max_uses_for_current_user()
                last_response_time_ms = session.get("last_response_time_ms")

                return render_template(
                    "dashboard.html",
                    text=text,
                    result=None,
                    suggestions=[],
                    logged_in=is_logged_in(),
                    user_email=session.get("user_email"),
                    remaining_uses=remaining_uses,
                    role=role,
                    usage_count=usage_count,
                    max_guest_uses=max_uses_local,
                    last_response_time_ms=last_response_time_ms,
                )

            if uploaded_file.filename.lower().endswith(".docx"):
                try:
                    doc = docx.Document(uploaded_file)
                    file_text = "\n".join(
                        [p.text for p in doc.paragraphs if p.text.strip()]
                    )
                    text = file_text.strip()
                    from_file = True
                except Exception as e:
                    flash(f"Error reading file: {e}")
            else:
                flash("Only .docx files are supported.")

        # 3) Run analysis if we have any text
        if text:
            # measure analysis time
            start_time = time.perf_counter()

            highlighted, suggestions_set = highlight_text(text)
            bias_pct, biased_count, total_words = calculate_bias_percentage(text)

            end_time = time.perf_counter()
            elapsed_ms = (end_time - start_time) * 1000.0

            # round for display and store in session
            session["last_response_time_ms"] = round(elapsed_ms, 1)

            suggestions = sorted(
                [{"word": w, "replacements": list(r)} for (w, r) in suggestions_set],
                key=lambda x: x["word"].lower()
            )

            result = {
                "highlighted": highlighted,
                "bias_pct": bias_pct,
                "biased_count": biased_count,
                "total_words": total_words,
                "from_file": from_file,
                "text": text,
            }

            # store last analysis for certificate + cleaned download
            session["last_analysis"] = {
                "bias_pct": bias_pct,
                "from_file": from_file,
                "total_words": total_words,
                "text": text,
                "suggestions": [
                    {"word": item["word"], "replacements": item["replacements"]}
                    for item in suggestions
                ],
            }

            if role != "subscriber":
                increment_usage()
        else:
            flash("Please paste text or upload a .docx file before scanning.")

    # 3) Values for bottom status bar and remaining uses
    remaining_uses = None
    if role != "subscriber":
        remaining_uses = max(0, max_uses - get_usage_count())

    usage_count = get_usage_count()
    max_uses = get_max_uses_for_current_user()
    last_response_time_ms = session.get("last_response_time_ms")

    return render_template(
        "dashboard.html",
        text=text,
        result=result,
        suggestions=suggestions,
        logged_in=is_logged_in(),
        user_email=session.get("user_email"),
        remaining_uses=remaining_uses,
        role=role,
        usage_count=usage_count,
        max_guest_uses=max_uses,
        last_response_time_ms=last_response_time_ms,
    )


# -------------------------------------------------
# Download cleaned document
# -------------------------------------------------
@app.route("/download_cleaned", methods=["POST"])
def download_cleaned():
    """
    Generate and download a cleaned .docx document based on the last analysis.
    Uses last_analysis['text'] and last_analysis['suggestions'].
    """
    last = session.get("last_analysis")
    if not last:
        flash("No recent analysis found. Please analyze text or upload a document first.")
        return redirect(url_for("dashboard"))

    original_text = last.get("text", "")
    suggestions_list = last.get("suggestions", [])

    if not original_text:
        flash("No original text available to clean.")
        return redirect(url_for("dashboard"))

    cleaned_text = apply_suggestions_to_text(original_text, suggestions_list)

    doc = docx.Document()
    for line in cleaned_text.split("\n"):
        doc.add_paragraph(line)

    bytes_io = io.BytesIO()
    doc.save(bytes_io)
    bytes_io.seek(0)

    return send_file(
        bytes_io,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="wordbalance_cleaned_document.docx",
    )


# -------------------------------------------------
# Certificate page
# -------------------------------------------------
@app.route("/certificate", methods=["GET", "POST"])
def certificate_page():
    """
    Certificate page:
      - requires login
      - requires last analysis bias_pct == 0 and from_file == True
      - generates PDF on POST
    """
    if not is_logged_in():
        flash("You must be logged in to generate a certificate.")
        return redirect(url_for("dashboard"))

    last = session.get("last_analysis")
    if not last:
        flash("No recent analysis found. Please analyze a document first.")
        return redirect(url_for("dashboard"))

    if last.get("bias_pct", 1) != 0 or not last.get("from_file", False):
        flash("Certificate can only be generated for uploaded documents with 0% bias.")
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        paper_title = request.form.get("paper_title", "").strip()
        if not paper_title:
            flash("Please enter a paper title.")
            return redirect(url_for("certificate_page"))

        pdf_bytes = generate_pdf_bytes(paper_title)

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{paper_title.replace(' ', '_')}_Certificate.pdf",
        )

    paper_title = request.args.get("title", "").strip()
    return render_template(
        "certificate.html",
        title=paper_title,
        logged_in=is_logged_in(),
        user_email=session.get("user_email"),
    )


# -------------------------------------------------
# Google login / logout
# -------------------------------------------------
@app.route("/login")
def login():
    redirect_uri = url_for("auth_callback", _external=True)
    print("Redirect URI being used:", redirect_uri)
    return google.authorize_redirect(redirect_uri)


@app.route("/auth/callback")
def auth_callback():
    """
    Google redirects back here after user approves/denies.
    We exchange the code for a token, read user info, and set session.
    """
    try:
        token = google.authorize_access_token()
        access_token = token.get("access_token")
        if not access_token:
            flash("Google login failed: no access token returned.")
            return redirect(url_for("dashboard"))

        userinfo_response = requests.get(
            "https://openidconnect.googleapis.com/v1/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}
        )
        if userinfo_response.status_code != 200:
            flash(f"Google login failed: userinfo error {userinfo_response.status_code}")
            return redirect(url_for("dashboard"))

        userinfo = userinfo_response.json()
    except Exception as e:
        flash(f"Google login failed: {e}")
        return redirect(url_for("dashboard"))

    if not userinfo:
        flash("Failed to retrieve user info from Google.")
        return redirect(url_for("dashboard"))

    email = userinfo.get("email")
    if not email:
        flash("No email address found in Google profile.")
        return redirect(url_for("dashboard"))

    # For now, treat all as normal 'user'; you can add subscriber logic here
    role = "user"

    session["user_email"] = email
    session["user_role"] = role
    session["usage_count"] = 0

    flash(f"Logged in as {email} (role: {role})")
    return redirect(url_for("dashboard"))


@app.route("/logout")
def logout():
    session.clear()
    flash("You have been logged out.")
    return redirect(url_for("dashboard"))


# -------------------------------------------------
# Error handler for file too large
# -------------------------------------------------
@app.errorhandler(RequestEntityTooLarge)
def handle_large_file(e):
    flash("Uploaded file is too large. Maximum allowed size is 50 MB.")
    return redirect(url_for("dashboard"))


# -------------------------------------------------
# Run
# -------------------------------------------------
